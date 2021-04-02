import requests
from lxml import etree
from docx import Document

#获取指定期数报纸下的文章url
def get_article_urls(url,cookie):
    headers = {
        "User-Agent": "Mozilla/5.0 (iPad; CPU OS 11_0 like Mac OS X) AppleWebKit/604.1.34 (KHTML, like Gecko) Version/11.0 Mobile/15A5341f Safari/604.1",
        "Host": "paper.i21st.cn",
        "Referer": "https://paper.i21st.cn/index_21je2.html",
        "cookie": cookie
        }
    try:
        response = requests.get(url, headers=headers)
        response.raise_for_status()
        response.encoding = "utf-8"
        print("已访问到{}，状态码：{}".format(url,response.status_code))
        text = response.text
        html = etree.HTML(text)
        urls = html.xpath("//div[@style='margin:0px 0px 7px 0px;']/a/@href")
        return urls
    except:
        return "获取URL异常"

#获取具体URL文章内容
def get_content(url,cookie):
    headers = {
        "User-Agent": "Mozilla/5.0 (iPad; CPU OS 11_0 like Mac OS X) AppleWebKit/604.1.34 (KHTML, like Gecko) Version/11.0 Mobile/15A5341f Safari/604.1",
        "Host": "paper.i21st.cn",
        "Referer": "https://paper.i21st.cn/index_21je2.html",
        "cookie": cookie
    }
    try:
        response = requests.get(url, headers=headers)
        response.raise_for_status()
        response.encoding = "utf-8"
        print("已访问到{}，状态码：{}".format(url, response.status_code))
        text = response.text
        html = etree.HTML(text)
        title = html.xpath("//h2[@class='title_txt']/text()")[0]
        paragraphs = html.xpath("//span[@class='i21stfanyi']/p/text()|//span[@class='i21stfanyi']/p/strong/text()")
        return title,paragraphs
    except:
        return "获取文章内容异常"

def get_cookie():
	with open("cookie.txt","r") as f:
		return f.read()




if __name__ == '__main__':
    cookie = get_cookie()
    document = Document()
    url = "https://paper.i21st.cn/index_21je1_issue_731.html"
    urls = get_article_urls(url,cookie)
    for i in urls:
        story_url = "https://paper.i21st.cn{}".format(i)
        title, paragraphs = get_content(story_url,cookie)
        print("正在访问到{}，请稍后^_^".format(title))
        document.add_heading(title, 0)
        for paragraph in paragraphs:
            p = document.add_paragraph(paragraph)
    document.save('test.docx')
    print("文件已保存")
    # url = "https://paper.i21st.cn/story/148824.html"
    # title,paragraphs = get_content(url,cookie)
    
    # print(paragraphs)
