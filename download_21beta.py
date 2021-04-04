import requests
from lxml import etree
from docx import Document
import os

#获取最新报纸期数
def get_newest_num(url):

    headers ={"User-Agent": "Mozilla/5.0 (iPad; CPU OS 11_0 like Mac OS X) AppleWebKit/604.1.34 (KHTML, like Gecko) Version/11.0 Mobile/15A5341f Safari/604.1",
            "Host": "paper.i21st.cn",
            "Referer": "https://paper.i21st.cn/index_21je2.html"
        }
    #调用get_url_text()获取URL的文本
    try:
        response = requests.get(url,headers=headers)
        response.raise_for_status()
        response.encoding = response.apparent_encoding
        print('正在获取响应码-------')
        print(response.status_code)#显示响应码
    #print(text)
        text = response.text
        html = etree.HTML(text)
        date = html.xpath("//div[@class='listPicDiv listPicDivM0']/text()")
    #print(date[0].strip())
        return "目前报纸最新为："+date[0].strip()
    except:
        return "产生异常"



def dl_path():
    #获取当前工作路径
    dir = os.getcwd()
    abs_dir = os.path.join(dir,"download")
    if os.path.exists(abs_dir):
        print('当前路径下/download已存在，正在执行下一步')
    else:
        print("当前路径下未发现/download文件夹，正在创建")
        os.mkdir(abs_dir)
    return abs_dir

#获取指定期数报纸下的文章url
def get_article_urls(url,cookie):
    headers = {
        "User-Agent": "Mozilla/5.0 (iPad; CPU OS 11_0 like Mac OS X) AppleWebKit/604.1.34 (KHTML, like Gecko) Version/11.0 Mobile/15A5341f Safari/604.1",
        "Host": "paper.i21st.cn",
        "Referer": "https://paper.i21st.cn/index_21je2.html",
        "cookie": cookie
        }
    try:
        response = requests.get(url, headers=headers)
        response.raise_for_status()
        response.encoding = "utf-8"
        print("已访问到{}，状态码：{}".format(url,response.status_code))
        text = response.text
        html = etree.HTML(text)
        urls = html.xpath("//div[@style='margin:0px 0px 7px 0px;']/a/@href")
        return urls
    except:
        return "获取URL异常"

#获取具体URL文章内容
def get_content(url,cookie):
    headers = {
        "User-Agent": "Mozilla/5.0 (iPad; CPU OS 11_0 like Mac OS X) AppleWebKit/604.1.34 (KHTML, like Gecko) Version/11.0 Mobile/15A5341f Safari/604.1",
        "Host": "paper.i21st.cn",
        "Referer": "https://paper.i21st.cn/index_21je2.html",
        "cookie": cookie
    }
    try:
        response = requests.get(url, headers=headers)
        response.raise_for_status()
        response.encoding = "utf-8"
        print("已访问到{}，状态码：{}".format(url, response.status_code))
        text = response.text
        html = etree.HTML(text)
        title = html.xpath("//h2[@class='title_txt']/text()")[0]
        paragraphs = html.xpath("//span[@class='i21stfanyi']/p/text()|//span[@class='i21stfanyi']/p/strong/text()")
        return title,paragraphs
    except:
        return "获取文章内容异常"

def get_cookie():
	with open("cookie.txt","r") as f:
		return f.read()




if __name__ == '__main__':

    while 1:
        grade = input("请输入要下载的年级，输入“q”退出：")
        if grade == 'q':
            break#无法退出
        if grade in ["1","2","3"]:
            break
        else:
            continue
    
    #获取该年级最新一期报纸，
    new_url = "https://paper.i21st.cn/index_21je{}.html".format(grade)
    print(get_newest_num(new_url))
    

    #获取用户所需哪几期报纸，保存在nums列表里
    while 1:
    #使用num保存用户所需下载的报纸期数
        num =input('输入需要下载报纸word版的期数，若批量下载以减号隔开；')
        try:
            if '-' in num:
                numlist = num.split('-')
                nums = list(range(int(numlist[0]),int(numlist[1])+1))
                break
            else:
                nums=[int(num)]
                break
        except:
            print("input error")
            continue




    cookie = get_cookie()
    dl = dl_path()
    for i in nums:
        url = "https://paper.i21st.cn/index_21je{}_issue_{}.html".format(grade,i)
        document = Document()
        #url = "https://paper.i21st.cn/index_21je1_issue_731.html"
        urls = get_article_urls(url,cookie)
        for j in urls:
            story_url = "https://paper.i21st.cn{}".format(j)
            title, paragraphs = get_content(story_url,cookie)
            print("正在访问到{}，请稍后^_^".format(title))
            document.add_heading(title, 0)
            for paragraph in paragraphs:
                p = document.add_paragraph(paragraph)
        document.save('{}/21je{}_issue_{}.docx'.format(dl,grade,i))
        print("21je{}_issue_{}文件已保存".format(grade,i))
    # url = "https://paper.i21st.cn/story/148824.html"
    # title,paragraphs = get_content(url,cookie)
    
    # print(paragraphs)
