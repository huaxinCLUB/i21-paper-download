import datetime
import re
from threading import  Thread
import PySide2
import requests
from PySide2 import QtCore
from PySide2.QtCore import QObject, Signal
from PySide2.QtGui import QIcon
from docx import Document
from lxml import etree
import os
from PySide2.QtUiTools import QUiLoader
from PySide2.QtWidgets import QApplication,QMessageBox,QFileDialog

# 信号库
class SignalStore(QObject):
    # 定义更改进度条的信号
    progress_update = Signal(int)
    # 定义更新日志的信号
    log_update = Signal(str)
    # 定义更新状态栏的信号
    statusbar_update = Signal(str)

# 实例化
so = SignalStore()


class Downloader():

    def __init__(self):
        #从文件中加载UI定义
        self.ui = QUiLoader().load("UI/main.ui")
        #设置软件图标icon
        self.ui.setWindowIcon(QIcon('images/logo_256.ico'))

        #给查询的按钮点击信号增加槽函数
        self.ui.checkButton.clicked.connect(self.checkButtonEvent)

        #给浏览按钮添加槽函数，选择下载路径
        self.ui.chooseDownloadPath.clicked.connect(self.downloadPath)
        #给下载按钮添加槽函数，根据用户输入的报纸期数和下载类型（PDF，word）进行下载
        self.ui.downloadButton.clicked.connect(self.action)

        #设置默认下载路径为：当前目录下的download文件夹
        cwd = os.getcwd()
        self.ui.dlPath.setText(os.path.join(cwd,"download"))

        #状态栏显示测试
        self.ui.statusbar.showMessage("今天是我陪伴你们的第 {} 天，今天也要加油鸭！".format(self.get_days()))

        # 连接信号到处理的slot函数
        so.progress_update.connect(self.setProgress)

        so.log_update.connect(self.setLog)

        so.statusbar_update.connect(self.setSuatus)


    #查询按钮的槽函数
    def checkButtonEvent(self):

        headers = {
            "User-Agent": "Mozilla/5.0 (iPad; CPU OS 11_0 like Mac OS X) AppleWebKit/604.1.34 (KHTML, like Gecko) Version/11.0 Mobile/15A5341f Safari/604.1",
            "Host": "paper.i21st.cn",
            "Referer": "https://paper.i21st.cn/index_21je2.html"
            }
        #获取年级的index，起始值为0
        grade = self.ui.grade.currentIndex()+1

        url = "https://paper.i21st.cn/index_21je{}.html".format(grade)
        #打印用户选择查询的年级及URL
        self.ui.logs.append("您正在查询的是{}，URL为：{}\n".format(self.ui.grade.currentText(),url))

        try:
            response = requests.get(url, headers=headers)
            response.raise_for_status()
            response.encoding = response.apparent_encoding
            self.ui.logs.append('正在获取响应码-------')
            #self.ui.logs.append 只能显示str类型格式
            self.ui.logs.append(str(response.status_code))  # 显示响应码
            # print(text)
            text = response.text
            html = etree.HTML(text)
            date = html.xpath("//div[@class='listPicDiv listPicDivM0']/text()")
            # print(date[0].strip())
            self.ui.logs.append("目前{}最新的报纸为：{}。请填写需要下载的报纸期数！" .format(self.ui.grade.currentText(),date[0].strip()))

            #将下载期数默认设置为目前最新的期数
            newest_num = eval(re.findall(r"第\d+期",date[0])[0][1:-1])
            # print(type(eval(newest_num)))
            self.ui.startNum.setValue(newest_num)
            self.ui.endNum.setValue(newest_num)

            #设置报纸期数的最大值为最新的期数
            self.ui.startNum.setRange(165,newest_num)
            self.ui.endNum.setRange(165,newest_num)

            #毒鸡汤上线
            so.statusbar_update.emit()


        except:
            QMessageBox.critical(self.ui,'错误','查询错误，请联系作者\n点击{}，提交issue'.format("<a href='https://github.com/huaxinCLUB/i21-paper-download/issues/new' style='color:blue'>这里</a>"))


    #浏览按钮的槽函数，点击浏览后弹窗选择下载路径，并在左侧文本框中显示下载路径
    def downloadPath(self):
        filePath = QFileDialog.getExistingDirectory(self.ui, "选择存储路径")
        self.ui.dlPath.setText(filePath)

        # 更新毒鸡汤
        so.statusbar_update.emit()







    #下载按钮的槽函数
    def action(self):
        #判断下载路径是否为空
        if self.ui.dlPath.text() == "":
            QMessageBox.critical(self.ui,'错误','还没有选择下载路径呢，急什么！')
        else:
            #判断选择的路径是否存在，如无则创建
            try:
                if os.path.exists(self.ui.dlPath.text()):
                    self.ui.logs.append('下载目录已存在，正在执行下一步\n*************************************\n')
                else:
                    self.ui.logs.append("未发现下载目录，正在创建**************\n*************************************\n")
                    os.mkdir(self.ui.dlPath.text())
            except:
                #若文件未正常创建，则弹窗提示用户重新输入
                QMessageBox.critical(self.ui, '错误', '路径有误，请重新选择下载路径！')

        #输出下载目录为self.ui.dlPath.text()
        self.ui.logs.append("文件保存目录为：{}\n*************************************\n".format(self.ui.dlPath.text()))

        #判断结束期数是否大于等于起始期数
        startNum = self.ui.startNum.value()
        endNum = self.ui.endNum.value()
        if startNum > endNum:
            QMessageBox.critical(self.ui, '错误', '小可爱，貌似结束期数应该大于等于起始期数吧！')
        #
        #输出下载报纸的年级（）和期数
        self.ui.logs.append("需要下载{}第 {} 期——第 {} 期 ， 共{}期报纸\n************************************\n".format(self.ui.grade.currentText(),startNum,endNum,endNum-startNum+1))

        #判断用户是否PDF和word都未选择
        # print(self.ui.pdfCheckBox.checkState()==PySide2.QtCore.Qt.CheckState.Unchecked)
        if self.ui.pdfCheckBox.checkState()==PySide2.QtCore.Qt.CheckState.Unchecked and self.ui.wordCheckBox.checkState()==PySide2.QtCore.Qt.CheckState.Unchecked:
            print(1 - self.ui.pdfCheckBox.isChecked() and self.ui.wordCheckBox.isChecked())
            QMessageBox.critical(self.ui, '错误', 'PDF和word总得爱一个吧')
            # 毒鸡汤上线
            so.statusbar_update.emit()

        #确定需要下载报纸的期数，保存在列表nums里
        nums = list(range(startNum, endNum + 1))

        #从文本框中获取cookie
        cookie = self.ui.cookie.toPlainText()
        self.ui.logs.append("已获取到的cookie为： {}\n".format(cookie))

        #判断是否要下载PDF格式的报纸
        if self.ui.pdfCheckBox.isChecked():
            for num in nums:
                pdf_url = "https://paper.i21st.cn/download/21je{}_{}.pdf".format(self.ui.grade.currentIndex()+1,num)
                self.ui.logs.append("正在下载第 {} 期PDF版报纸，URL为：{}  请稍后****".format(num,pdf_url))
                #创建下载PDF的线程
                down_pdf_thread = Thread(target=self.down_pdf, args=(pdf_url,cookie,num))
                down_pdf_thread.start()
                # self.down_pdf(pdf_url,cookie,num)
                # 毒鸡汤上线
                so.statusbar_update.emit()

        #判断是否下载docx格式的报纸
        if self.ui.wordCheckBox.isChecked():
            for doc_num in nums:
                doc_url = "https://paper.i21st.cn/index_21je{}_issue_{}.html".format(self.ui.grade.currentIndex()+1,doc_num)
                #创建下载docx的线程
                down_docx_thread = Thread(target=self.save_docx, args=(doc_url,cookie,doc_num))
                down_docx_thread.start()
                # self.save_docx(doc_url,cookie,doc_num)
                # url = "https://paper.i21st.cn/index_21je1_issue_731.html"
                # 毒鸡汤上线
                so.statusbar_update.emit()




    #下载指定报纸的PDF版
    def down_pdf(self,url,cookie,num):
        headers = {
            "User-Agent": "Mozilla/5.0 (iPad; CPU OS 11_0 like Mac OS X) AppleWebKit/604.1.34 (KHTML, like Gecko) Version/11.0 Mobile/15A5341f Safari/604.1",
            "Host": "paper.i21st.cn",
            "Referer": "https://paper.i21st.cn/index_21je2.html",
            "cookie": cookie}
        try:
            r = requests.get(url, headers=headers, stream=True)
            # 获取pdf文件大小
            file_size_str = r.headers['Content-Length']
        except:
            QMessageBox.warning(self.ui,'获取文件大小失败','请重新获取cookie')
        # 将文件大小转化成MB单位
        file_size = int(file_size_str) /1024
        #设置进度条最大值为文件大小
        self.ui.progressBar.setRange(1,file_size)

        filename = "21je{}_{}.pdf".format(self.ui.grade.currentIndex()+1, num)
        dl_path_name = os.path.join(self.ui.dlPath.text(), filename)
        with open(dl_path_name, 'wb') as f:
            #将日志信号传递给主线程
            so.log_update.emit("当前正在下载{}".format(filename))
            # self.ui.logs.append("当前正在下载{}".format(filename))
            count = 1
            for chunk in r.iter_content(chunk_size=1024):
                if chunk:
                    f.write(chunk)
                    # 获取当前下载进度的百分比
                    so.progress_update.emit(count)
                    # self.ui.progressBar.setValue(count)
                    # percent = count / 1024 / file_size
                    # self.ui.logs.append("\r", "{:.2%}{}".format(percent, "#" * int(percent * 72)), end="",flush=True)  # 100.0%加上72个#正好占满一行
                    # self.ui.logs.append("当前正在下载{}，已下载{}MB /{}MB ".format(filename,count/512,file_size))
                    count = count + 1
            self.ui.logs.append("**************{}已下载完毕*****************\n".format(filename))
            so.progress_update.emit(0)

    # 获取指定期数报纸下的文章url
    def get_article_urls(self, url, cookie):
        headers = {
            "User-Agent": "Mozilla/5.0 (iPad; CPU OS 11_0 like Mac OS X) AppleWebKit/604.1.34 (KHTML, like Gecko) Version/11.0 Mobile/15A5341f Safari/604.1",
            "Host": "paper.i21st.cn",
            "Referer": "https://paper.i21st.cn/index_21je2.html",
            "cookie": cookie
        }
        try:
            response = requests.get(url, headers=headers)
            response.raise_for_status()
            response.encoding = "utf-8"
            so.log_update.emit("已访问到{}，状态码：{}".format(url, response.status_code))
            text = response.text
            html = etree.HTML(text)
            urls = html.xpath("//div[@style='margin:0px 0px 7px 0px;']/a/@href")
            return urls
        except:
            return "获取URL异常"

    # 获取具体URL文章内容
    def get_content(self, url, cookie):
        headers = {
            "User-Agent": "Mozilla/5.0 (iPad; CPU OS 11_0 like Mac OS X) AppleWebKit/604.1.34 (KHTML, like Gecko) Version/11.0 Mobile/15A5341f Safari/604.1",
            "Host": "paper.i21st.cn",
            "Referer": "https://paper.i21st.cn/index_21je2.html",
            "cookie": cookie
        }
        try:
            response = requests.get(url, headers=headers)
            response.raise_for_status()
            response.encoding = "utf-8"
            # print("已访问到{}，状态码：{}".format(url, response.status_code))
            text = response.text
            html = etree.HTML(text)
            title = html.xpath("//h2[@class='title_txt']/text()")[0]
            paragraphs = html.xpath("//span[@class='i21stfanyi']/p/text()|//span[@class='i21stfanyi']/p/strong/text()")
            return title, paragraphs
        except:
            QMessageBox.critical(self.ui,'错误','获取文章内容失败，请联系作者\n点击{}，提交issue'.format("<a href='https://github.com/huaxinCLUB/i21-paper-download/issues/new' style='color:blue'>这里</a>"))

   #保存docx文件
    def save_docx(self,doc_url,cookie,doc_num):
        document = Document()
        urls = self.get_article_urls(doc_url, cookie)
        self.ui.progressBar.setRange(1,len(urls))
        count = 1
        for j in urls:
            story_url = "https://paper.i21st.cn{}".format(j)
            title, paragraphs = self.get_content(story_url, cookie)
            # print("正在访问到{}，请稍后^_^".format(title))
            document.add_heading(title, 0)
            for paragraph in paragraphs:
                p = document.add_paragraph(paragraph)

            so.progress_update.emit(count)
            count += 1

        docx_name = os.path.join(self.ui.dlPath.text(), "21je{}_issue_{}.docx".format(self.ui.grade.currentIndex()+1,doc_num))

        document.save(docx_name)
        so.log_update.emit("21je{}_issue_{}文件已保存".format(self.ui.grade.currentIndex()+1, doc_num))


    # 处理进度的slot函数
    def setProgress(self,value):
        self.ui.progressBar.setValue(value)

    def setLog(self, text):
        self.ui.logs.append(text)

    def setSuatus(self):
        self.ui.statusbar.showMessage(self.soup())

    #毒鸡汤函数
    def soup(self):
        try:
            #接口一
            url = 'https://v1.alapi.cn/api/soul'
            res = requests.get(url)
            res.encoding = res.apparent_encoding
            return res.json()['data']['title']
        except:
            try:
                #接口二
                url = "https://soul-soup.fe.workers.dev/"
                res = requests.get(url)
                res.encoding = res.apparent_encoding
                return res.json()['title']
            except:
                return "接口异常，请联系作者"

    #获取陪伴天数
    def get_days(self):
        #获取今天的时间
        new_date = datetime.datetime.now()  # 现在时间
        data_str = new_date.strftime('%Y-%m-%d')  # 格式化时间
        oneDay = datetime.datetime(2021, 4, 19)
        difference = new_date.toordinal() - oneDay.toordinal()
        return difference


QtCore.QCoreApplication.setAttribute(QtCore.Qt.AA_ShareOpenGLContexts)
app = QApplication([])
downloader = Downloader()
downloader.ui.show()
app.exec_()